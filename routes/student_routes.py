import json
import re
import os
import shutil
import pdfplumber
from docx import Document
from fastapi import APIRouter, Form, File, UploadFile, Depends, HTTPException
from dependencies import get_current_student

from database import (
    get_all_question_papers,
    get_answer_scheme_by_paper,
    add_submission,
    get_student_analytics
)

from nlp.similarity import calculate_marks, generate_explanation

router = APIRouter(prefix="/student", tags=["Student"])


# ============================================
# GET ALL PAPERS
# ============================================
@router.get("/papers")
def get_papers():
    return get_all_question_papers()


# ============================================
# GET QUESTIONS FROM ANSWER SCHEME
# ============================================
@router.get("/paper/{paper_id}/questions")
def get_questions(paper_id: int):

    scheme = get_answer_scheme_by_paper(paper_id)
    if not scheme:
        return {"error": "Answer scheme not found"}

    scheme_path = scheme[2]

    with open(scheme_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    pattern = r"Q(\d+)\|(\d+):(.*?)(?=Q\d+\|\d+:|$)"
    matches = re.findall(pattern, content, re.DOTALL)

    questions = []
    for q_no, marks, _ in matches:
        questions.append({
            "question_number": int(q_no),
            "marks": int(marks)
        })

    return questions


# ============================================
# SUBMIT PAPER
# ============================================
@router.post("/submit-paper")
async def submit_paper(
    paper_id: int = Form(...),
    answers: str = Form(None),
    file: UploadFile = File(None),
    student_id: int = Form(None), # Kept for backward compat with frontend FormData but ignored
    current_user: dict = Depends(get_current_student)
):
    student_id = current_user["id"]

    # Setup permanent storage
    STUDENT_SUBMISSION_FOLDER = "uploads/student_submissions"
    os.makedirs(STUDENT_SUBMISSION_FOLDER, exist_ok=True)
    stored_file_path = None

    # ---------------------------------------
    # VALIDATE PAPER
    # ---------------------------------------
    papers = get_all_question_papers()
    paper_ids = [p[0] for p in papers]

    if paper_id not in paper_ids:
        raise HTTPException(status_code=400, detail="Invalid paper selected.")

    # ---------------------------------------
    # LOAD & EXTRACT ANSWER SCHEME
    # ---------------------------------------
    scheme = get_answer_scheme_by_paper(paper_id)
    if not scheme:
        raise HTTPException(status_code=404, detail="Answer scheme not found")

    scheme_path = scheme[2]
    scheme_content = ""

    # Support PDF, DOCX in Answer Scheme too!
    if scheme_path.lower().endswith(".pdf"):
        from ocr import extract_text_from_pdf
        scheme_content = extract_text_from_pdf(scheme_path)
    elif scheme_path.lower().endswith(".docx"):
        from io import BytesIO
        with open(scheme_path, "rb") as f:
            doc = Document(BytesIO(f.read()))
            scheme_content = "\n".join([p.text for p in doc.paragraphs])
    else:
        # Fallback to plain text
        with open(scheme_path, "r", encoding="utf-8", errors="ignore") as f:
            scheme_content = f.read()

    # Log teacher content for debugging
    print("\n================ TEACHER SCHEME ================")
    print(scheme_content[:1000] + "...") 
    print("================================================\n")

    # 1. Try Strict Pattern: Q1|10:
    pattern_strict = r"Q\s*(\d+)\s*\|\s*(\d+)\s*[:\-\._\s](.*?)(?=Q\s*\d+\s*\|\s*\d+\s*[:\-\._\s]|$)"
    matches = re.findall(pattern_strict, scheme_content, re.DOTALL | re.IGNORECASE)

    if not matches:
        # 2. Try Semi-Strict: 1|10:
        pattern_semi = r"(?:^|\s)(\d+)\s*\|\s*(\d+)\s*[:\-\._\s](.*?)(?=(?:^|\s)\d+\s*\|\s*\d+\s*[:\-\._\s]|$)"
        matches = re.findall(pattern_semi, scheme_content, re.DOTALL)

    if not matches:
        # 3. Try Lazy (OCR Friendly): 1. or 1: or 1_ or 1)
        # We look for a number (1-3 digits) followed by a delimiter
        pattern_lazy = r"(?:^|\s)([0-9]{1,3})[\.\:\-\)_\|]\s*(.*?)(?=(?:^|\s)[0-9]{1,3}[\.\:\-\)_\|]|$)"
        found = re.findall(pattern_lazy, scheme_content, re.DOTALL)
        matches = [(m[0], "1", m[1]) for m in found] # Default to 1 mark

    print(f"DEBUG: Found {len(matches)} questions in teacher scheme.")
    for m in matches[:3]: print(f"  - Q{m[0]} ({m[1]} marks): {m[2][:50]}...")

    student_answers = {}

    # ---------------------------------------
    # CASE 1: TYPED ANSWERS
    # ---------------------------------------
    if answers:
        student_answers = json.loads(answers)

    # ---------------------------------------
    # CASE 2: FILE UPLOAD
    # ---------------------------------------
    elif file:

        extracted_text = ""

        from io import BytesIO
        file_content = await file.read()

        # PDF → OCR
        if file.filename.endswith(".pdf"):

            import tempfile
            from ocr import extract_text_from_pdf

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp:
                temp.write(file_content)
                temp_path = temp.name

            extracted_text = extract_text_from_pdf(temp_path)

        # DOCX
        elif file.filename.lower().endswith(".docx"):
            from io import BytesIO
            doc = Document(BytesIO(file_content))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"

        # IMAGES (JPG, PNG)
        elif file.filename.lower().endswith((".jpg", ".jpeg", ".png")):
            from ocr import extract_text_from_image
            import numpy as np
            import cv2

            # Read image directly from memory
            nparr = np.frombuffer(file_content, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            
            if img is not None:
                extracted_text = extract_text_from_image(img)
            else:
                raise HTTPException(status_code=400, detail="Invalid image file.")

        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please use PDF, DOCX, JPG, or PNG.")

        # Save the file permanently
        stored_file_path = os.path.join(STUDENT_SUBMISSION_FOLDER, f"S{student_id}_P{paper_id}_{file.filename}")
        with open(stored_file_path, "wb") as buffer:
            buffer.write(file_content)

        # ---------------------------------------
        # DEBUG: SHOW OCR TEXT
        # ---------------------------------------
        print("\n================ OCR OUTPUT ================")
        print(extracted_text)
        print("============================================\n")

        # Fix common OCR mistakes
        extracted_text = extracted_text.replace("Ql", "Q1")
        extracted_text = extracted_text.replace("QI", "Q1")
        extracted_text = extracted_text.replace("Qi", "Q1")
        extracted_text = extracted_text.replace("Q l", "Q1")
        extracted_text = extracted_text.replace("Q I", "Q1")
        extracted_text = extracted_text.replace("QS", "Q5")
        extracted_text = extracted_text.replace("Qs", "Q5")
        extracted_text = extracted_text.replace("Q S", "Q5")
        extracted_text = extracted_text.replace("Q s", "Q5")
        extracted_text = extracted_text.replace("QO", "Q0")
        extracted_text = extracted_text.replace("Qo", "Q0")
        extracted_text = extracted_text.replace("Q O", "Q0")
        extracted_text = extracted_text.replace("Q o", "Q0")

        # ---------------------------------------
        # OCR FRIENDLY QUESTION SPLIT (More flexible)
        # ---------------------------------------
        split_pattern = r"(?:^|\n)\s*(?:Q|Question|Ans|Answer|Q\s*)?\.?\s*([0-9]+)[\:\.\-\)\;]\s*(.*?)(?=(?:^|\n)\s*(?:Q|Question|Ans|Answer|Q\s*)?\.?\s*[0-9]+[\:\.\-\)\;]|$)"

        found_answers = re.findall(split_pattern, extracted_text, re.DOTALL | re.IGNORECASE)
        student_answers = {m[0]: m[1].strip() for m in found_answers}

        print(f"DEBUG: Extracted {len(student_answers)} answers from student OCR.")
        for q, a in list(student_answers.items())[:10]: print(f"  - Q{q}: {a[:50]}...")

    else:
        raise HTTPException(status_code=400, detail="No answers provided.")

    # ---------------------------------------
    # EVALUATE PAPER
    # ---------------------------------------
    total_marks = 0
    total_obtained = 0
    detailed_results = []
    from datetime import datetime
    submission_time = datetime.utcnow().isoformat()

    # If OCR text exists but answers weren't found via regex, we prepare for semantic search
    from nlp.similarity import model as emb_model
    from sklearn.metrics.pairwise import cosine_similarity
    
    ocr_sentences = []
    ocr_embeddings = []
    
    if file and extracted_text:
        import nltk
        ocr_sentences = nltk.sent_tokenize(extracted_text)
        if ocr_sentences:
            ocr_embeddings = emb_model.encode(ocr_sentences)

    for q_no, marks, model_answer in matches:

        q_no = int(q_no)
        marks = int(marks)
        total_marks += marks

        student_answer = student_answers.get(str(q_no), "")
        
        # ---------------------------------------
        # FALLBACK: SEMANTIC SEARCH
        # If student answer is empty but we have OCR text
        # ---------------------------------------
        if not student_answer.strip() and ocr_embeddings is not None and len(ocr_embeddings) > 0:
            best_chunk = ""
            best_sim = 0
            
            model_emb = emb_model.encode([model_answer])
            
            # Use a slightly larger window context (combine current sentence with next)
            for i in range(len(ocr_sentences)):
                window_text = ocr_sentences[i]
                if i + 1 < len(ocr_sentences):
                    window_text += " " + ocr_sentences[i+1]
                if i + 2 < len(ocr_sentences):
                    window_text += " " + ocr_sentences[i+2]
                    
                window_emb = emb_model.encode([window_text])
                sim = cosine_similarity(model_emb, window_emb)[0][0]
                
                if sim > best_sim:
                    best_sim = sim
                    best_chunk = window_text
            
            # If the best chunk has reasonable relevance, assume it's the answer
            if best_sim > 0.3:
                student_answer = best_chunk
                print(f"Fallback matched Q{q_no} with similarity {best_sim:.2f}")

        obtained, breakdown = calculate_marks(
            model_answer.strip(),
            student_answer,
            marks
        )

        explanation = generate_explanation(breakdown)

        total_obtained += float(obtained)

        detailed_results.append({
            "question": int(q_no),
            "marks_awarded": float(obtained),
            "max_marks": int(marks),
            "semantic_score": float(breakdown["semantic"]),
            "keyword_score": float(breakdown["keyword"]),
            "length_score": float(breakdown["length"]),
            "concept_coverage": breakdown["concept_coverage"],
            "explanation": explanation
        })

        add_submission(
            student_id,
            paper_id,
            q_no,
            student_answer,
            float(breakdown["semantic"]),
            float(obtained),
            stored_file_path,
            submission_time
        )

    return {
        "total_marks": int(total_marks),
        "total_obtained": float(total_obtained),
        "details": detailed_results
    }


# ============================================
# GET PAPER DETAILS
# ============================================
@router.get("/paper/{paper_id}")
def get_paper_details(paper_id: int):

    papers = get_all_question_papers()

    for paper in papers:
        if paper[0] == paper_id:

            scheme = get_answer_scheme_by_paper(paper_id)

            return {
                "id": paper[0],
                "subject": paper[1],
                "exam_title": paper[2],
                "total_marks": paper[3],
                "file_path": paper[4],
                "is_ready": True if scheme else False
            }

    return {"error": "Paper not found"}


# ============================================
# GET STUDENT ANALYTICS
# ============================================
@router.get("/analytics/{student_id}")
def get_analytics(student_id: int, current_user: dict = Depends(get_current_student)):
    if student_id != current_user["id"]:
        return {"error": "Unauthorized"}
    history = get_student_analytics(student_id)
    return {"history": history}